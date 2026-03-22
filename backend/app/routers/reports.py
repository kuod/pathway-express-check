"""
Report generation endpoints.
Supports HTML (printable preprint), DOCX, and CSV/Excel data exports.
"""

import base64
import io
import zipfile
from datetime import date

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt
from fastapi import APIRouter
from fastapi.responses import Response, StreamingResponse
from pathlib import Path
from jinja2 import Environment, FileSystemLoader, select_autoescape

from app.models.schemas import ReportRequest
from app.services.analysis import build_dataframe

router = APIRouter()

_JINJA_ENV = Environment(
    loader=FileSystemLoader(str(Path(__file__).parent.parent / "templates")),
    autoescape=select_autoescape(["html"]),
)


# ---------------------------------------------------------------------------
# Figure helpers
# ---------------------------------------------------------------------------

def _expression_heatmap_png(expression_matrix: dict, figsize=(14, 8)) -> bytes:
    df = build_dataframe(expression_matrix)
    log_df = np.log2(df + 1)
    fig, ax = plt.subplots(figsize=figsize)
    sns.heatmap(
        log_df,
        cmap="viridis",
        ax=ax,
        cbar_kws={"label": "log₂(TPM+1)"},
        xticklabels=True,
        yticklabels=True,
    )
    ax.set_title("Median Gene Expression Across GTEx Tissues")
    ax.set_xlabel("Tissue")
    ax.set_ylabel("Gene")
    plt.xticks(rotation=90, fontsize=6)
    plt.yticks(fontsize=7)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    return buf.getvalue()


def _correlation_heatmap_png(correlation_matrix: dict, figsize=(10, 8)) -> bytes:
    corr_df = pd.DataFrame(correlation_matrix)
    fig, ax = plt.subplots(figsize=figsize)
    mask = np.triu(np.ones_like(corr_df, dtype=bool))
    sns.heatmap(
        corr_df,
        mask=mask,
        cmap="RdBu_r",
        center=0,
        vmin=-1,
        vmax=1,
        ax=ax,
        cbar_kws={"label": "Pearson r"},
        square=True,
        xticklabels=True,
        yticklabels=True,
    )
    ax.set_title("Gene Co-expression Correlation Matrix")
    plt.xticks(rotation=45, ha="right", fontsize=7)
    plt.yticks(fontsize=7)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    return buf.getvalue()


def _pca_scatter_png(pca_coords: dict, figsize=(7, 6)) -> bytes:
    if not pca_coords:
        return b""
    genes = list(pca_coords.keys())
    xs = [pca_coords[g][0] for g in genes]
    ys = [pca_coords[g][1] for g in genes]
    fig, ax = plt.subplots(figsize=figsize)
    ax.scatter(xs, ys, s=60, alpha=0.8)
    for gene, x, y in zip(genes, xs, ys):
        ax.annotate(gene, (x, y), textcoords="offset points", xytext=(4, 4), fontsize=7)
    ax.set_xlabel("PC1")
    ax.set_ylabel("PC2")
    ax.set_title("PCA of Gene Expression Across Tissues")
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------

@router.post("/html", response_class=Response)
def generate_html_report(req: ReportRequest):
    """Generate a printable HTML preprint-style report."""
    expr_png = _expression_heatmap_png(req.expression.expression_matrix)
    corr_png = _correlation_heatmap_png(req.coexpression.correlation_matrix)
    pca_png = _pca_scatter_png(req.coexpression.pca_coordinates)

    stats_df = pd.DataFrame(req.coexpression.summary_stats).T
    stats_html = stats_df.to_html(classes="stats-table", border=0, float_format="{:.3f}".format)

    tmpl = _JINJA_ENV.get_template("report.html")
    html = tmpl.render(
        pathway_name=req.pathway_name,
        analysis_date=date.today().isoformat(),
        dataset_info=req.dataset_info,
        genes_found=req.expression.genes_found,
        genes_not_found=req.expression.genes_not_found,
        consistency_score=round(req.coexpression.pathway_consistency_score, 4),
        top_pairs=req.coexpression.top_correlated_pairs[:10],
        stats_table=stats_html,
        expr_img=base64.b64encode(expr_png).decode(),
        corr_img=base64.b64encode(corr_png).decode(),
        pca_img=base64.b64encode(pca_png).decode() if pca_png else "",
        pca_variance=req.coexpression.pca_variance_explained,
        include_methods=req.include_methods,
    )
    return Response(content=html, media_type="text/html")


@router.post("/docx")
def generate_docx_report(req: ReportRequest):
    """Generate a Word document preprint template."""
    doc = Document()

    # Title
    doc.add_heading(f"Co-expression Analysis: {req.pathway_name}", 0)
    doc.add_paragraph(f"Date: {date.today().isoformat()}  |  Dataset: {req.dataset_info}")

    # Abstract
    doc.add_heading("Abstract", 1)
    n_found = len(req.expression.genes_found)
    n_total = len(req.pathway_genes)
    score = round(req.coexpression.pathway_consistency_score, 3)
    doc.add_paragraph(
        f"We assessed the co-expression of {n_total} genes from the {req.pathway_name} pathway "
        f"using median RNA-seq expression data from {req.dataset_info}. "
        f"{n_found} of {n_total} genes were found in the dataset. "
        f"The pathway coherence score (mean pairwise Pearson correlation on log₂(TPM+1) values) "
        f"was {score}. [Expand with key findings.]"
    )

    # Introduction placeholder
    doc.add_heading("Introduction", 1)
    doc.add_paragraph("[Introduce the biological context of the pathway and the rationale for co-expression analysis.]")

    # Methods
    if req.include_methods:
        doc.add_heading("Methods", 1)
        doc.add_heading("Gene Sets", 2)
        doc.add_paragraph(
            "Gene sets were obtained from the Molecular Signatures Database (MSigDB) via gseapy. "
            f"The {req.pathway_name} gene set was used as input."
        )
        doc.add_heading("Expression Data", 2)
        doc.add_paragraph(
            f"Median transcript per million (TPM) expression values were retrieved from the "
            f"{req.dataset_info} via the GTEx Portal REST API (gtexportal.org/api/v2). "
            "Only protein-coding genes with versioned GENCODE v26 identifiers were included."
        )
        doc.add_heading("Co-expression Analysis", 2)
        doc.add_paragraph(
            "Expression values were log₂(TPM+1) transformed. "
            "Pairwise Pearson correlations were computed across all GTEx tissues. "
            "Pathway coherence was summarized as the mean of all upper-triangle correlation values. "
            "Principal component analysis (PCA) was performed on the log-transformed expression matrix "
            "with genes as observations and tissues as features."
        )

    # Results
    doc.add_heading("Results", 1)
    doc.add_paragraph(
        f"Of {n_total} pathway genes, {n_found} were found in {req.dataset_info}. "
        f"The overall pathway consistency score was {score}."
    )
    if req.expression.genes_not_found:
        doc.add_paragraph(
            f"Genes not found: {', '.join(req.expression.genes_not_found)}."
        )
    doc.add_paragraph("[Insert expression heatmap and correlation heatmap figures here.]")
    doc.add_paragraph("[Describe top co-expressed pairs and tissue-specific expression patterns.]")

    # Top pairs table
    doc.add_heading("Top Co-expressed Gene Pairs", 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Gene A", "Gene B", "Pearson r"
    for pair in req.coexpression.top_correlated_pairs[:10]:
        row = table.add_row().cells
        row[0].text = pair["gene_a"]
        row[1].text = pair["gene_b"]
        row[2].text = str(pair["pearson_r"])

    # Discussion & References
    doc.add_heading("Discussion", 1)
    doc.add_paragraph("[Interpret the co-expression patterns in the context of pathway biology.]")
    doc.add_heading("References", 1)
    doc.add_paragraph(
        "GTEx Consortium. The GTEx Consortium atlas of genetic regulatory effects across human tissues. "
        "Science 369, 1318–1330 (2020)."
    )
    doc.add_paragraph(
        "Liberzon A, et al. The Molecular Signatures Database Hallmark Gene Set Collection. "
        "Cell Syst 1, 417–425 (2015)."
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = req.pathway_name.replace(" ", "_") + "_coexpression_report.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/figures/zip")
def download_figures(req: ReportRequest):
    """Download all figures as a ZIP archive of high-res PNGs."""
    expr_png = _expression_heatmap_png(req.expression.expression_matrix, figsize=(18, 10))
    corr_png = _correlation_heatmap_png(req.coexpression.correlation_matrix, figsize=(12, 10))
    pca_png = _pca_scatter_png(req.coexpression.pca_coordinates)

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("expression_heatmap.png", expr_png)
        zf.writestr("correlation_heatmap.png", corr_png)
        if pca_png:
            zf.writestr("pca_scatter.png", pca_png)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": 'attachment; filename="pathway_figures.zip"'},
    )


@router.post("/data/csv")
def download_csv(req: ReportRequest):
    """Download expression matrix and summary stats as a multi-sheet Excel file."""
    expr_df = build_dataframe(req.expression.expression_matrix)
    stats_df = pd.DataFrame(req.coexpression.summary_stats).T

    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        expr_df.to_excel(writer, sheet_name="Expression_TPM")
        np.log2(expr_df + 1).to_excel(writer, sheet_name="Expression_log2TPM")
        stats_df.to_excel(writer, sheet_name="Summary_Stats")
        corr_df = pd.DataFrame(req.coexpression.correlation_matrix)
        corr_df.to_excel(writer, sheet_name="Correlation_Matrix")
    buf.seek(0)
    filename = req.pathway_name.replace(" ", "_") + "_expression_data.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
